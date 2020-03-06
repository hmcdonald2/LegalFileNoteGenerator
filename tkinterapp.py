#https://likegeeks.com/python-gui-examples-tkinter-tutorial/
#http://effbot.org/tkinterbook/label.htm

from tkinter import *

from tkinter.ttk import *

from tkinter import Menu

from tkinter import filedialog

import docx

from os import path

window = Tk()

window.title("Create a new File Note")

window.geometry('500x800')

window.configure(bg='gray28')

#functions

def newfn():
    #takes an existing word template file note and renames to current date and time
    import docx
    from time import localtime, time, asctime
    import os
    from datetime import date
    # os.chdir('INTENDED DESTINATION OF FILE NOTE')
    today = date.today()
    largetime = asctime(localtime(time()))
    smalltime = largetime[11:13] + "." + largetime[14:16]
    matter = mtrc.get()
    d1 = docx.Document()
    d2 = d1.save(f'File Note {matter} {today} {smalltime}.docx')
    d2 = docx.Document(f"File Note {matter} {today} {smalltime}.docx")
    d2.add_paragraph(f"Client:   {matter}")
    d2.add_paragraph(content.get(1.0,END))
    d2.save(f'File Note {matter} {today} {smalltime}.docx')

#todo - add variables for multiple choices re attendance type and attendance with


#add menu

menu = Menu(window)

new_item = Menu(menu)

new_item.add_command(label='New')

menu.add_cascade(label='File', menu=new_item)

window.config(menu=menu)


#file = filedialog.askopenfilename()

#files = filedialog.askopenfilenames()

#file = filedialog.askopenfilename(initialdir= path.dirname(__file__))


#add label / description to the top of the screen

lbl = Label(window, text="App to be used to create new file notes quickly", font=("Arial Bold", 15))

lbl.grid(pady=10)

#add label / instructions

lbl2 = Label(window, text="Select the matter name", font=("Arial Bold", 15))

lbl2.grid(column=0, row=5, pady=10)

#add text box to be entered

mtr = Entry(window, width=25)

mtr.grid(column=0, row=10, pady=10)

mtr.focus

#add multiple choice box to be clicked

mtrc = Combobox(window)

mtrc['values']=('Johnson', 'Jackson', 'Fred', 'Sully')

mtrc.current(1)

mtrc.grid(column=0, row=15)


#add label / instructions

lbl2 = Label(window, text="Attendance", font=("Arial Bold", 15))

lbl2.grid(column=0, row=20, pady=10)

#add multiple choice box to be clicked

AttendType = Combobox(window)

AttendType['values']=('Call in', 'Call out', 'Conference', 'Other' )

AttendType.current(1)

AttendType.grid(column=0, row=25, pady=10)

#add label / instructions

lbl2 = Label(window, text="In attendance with", font=("Arial Bold", 15))

lbl2.grid(column=0, row=26, pady=10)

#add multiple choice box to be clicked

AttendOn = Combobox(window)

AttendOn['values']=('Our client', 'Other side', 'Counsel', 'Other' )

AttendOn.current(1)

AttendOn.grid(column=0, row=27, pady=10)

#add label / instructions

lbl2 = Label(window, text="Enter content of the file note", font=("Arial Bold", 15))

lbl2.grid(column=0, row=30, pady=10)

#ADD text field to be the content of the file note

content = Text(window, width=25, bg='white smoke')

content.grid(column=0, row=35, pady=10)


#add button which is currently set to print matter selected from combobox

# def clicked():

    # ans = 'The matter name is ' + mtrc.get()
    #
    # lbl.configure(text=ans)


lbl.grid(column=0, row=40, pady=10)

btn = Button(window, text='Create new file note', command=newfn)

btn.grid(column=0, row=45, pady=10)


window.mainloop()

